from docx import Document
from datetime import datetime, timedelta
import os   # <-- added

def generate_prescription_file(patient_name, symptoms, duration, allergy, diagnosis, medicine, dosage, med_days):
    doc = Document()

    doc.add_heading("Medical Prescription", level=1)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    # Patient Details
    doc.add_heading("Patient Details", level=2)
    doc.add_paragraph(f"Name: {patient_name}")
    doc.add_paragraph(f"Symptoms: {symptoms}")
    doc.add_paragraph(f"Duration of illness: {duration} days")
    doc.add_paragraph(f"Allergies: {allergy if allergy else 'No known allergies'}")

    # Diagnosis
    doc.add_heading("Diagnosis", level=2)
    doc.add_paragraph(diagnosis)

    # Medication
    doc.add_heading("Medication", level=2)
    doc.add_paragraph(f"Medicine: {medicine}")
    doc.add_paragraph(f"Dosage Instructions: {dosage}")
    doc.add_paragraph(f"Medication Duration: {med_days} days")

    # Next follow-up
    next_follow_up = (datetime.now() + timedelta(days=med_days)).strftime('%d-%m-%Y')
    doc.add_paragraph(f"Next Follow-Up Date: {next_follow_up}")

    # General Advice
    doc.add_heading("General Advice", level=2)
    doc.add_paragraph(
        "• Take medication as prescribed.\n"
        "• Maintain hydration.\n"
        "• Take adequate rest.\n"
        "• Monitor symptoms."
    )

    # Save to Downloads folder
    download_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    os.makedirs(download_dir, exist_ok=True)

    file_name = os.path.join(
        download_dir,
        f"Prescription_{patient_name.replace(' ', '_')}.docx"
    )

    doc.save(file_name)

    return file_name
